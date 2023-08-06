import os
import re
import sys
from configparser import ConfigParser
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl import load_workbook
from colorama import init, Fore


class Project(object):
    p_comname = ''  # 公司名称
    p_order = ''  # 序号
    p_name = ''  # 项目名
    p_start = ''
    p_end = ''
    p_cost = ''
    p_people = ''  # 人数
    p_owner = ''  # 项目负责人
    p_rnd = ''  # 研发人员
    p_money = ''  # 总预算


def check_and_change(doc, replace):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
    （key 和 value 都是replace_dict中的键值对。）
    """
    for para in doc.paragraphs:
        for i in range(len(para.runs)):
            # print(">>>" + para.runs[i].text)
            for key, value in replace.items():
                if key in para.runs[i].text:
                    print(key + "-->" + value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return doc


def replace_tables(doc, replace):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for i in range(len(para.runs)):
                        # print(">>>" + para.runs[i].text)
                        for key, value in replace.items():
                            if key in para.runs[i].text:
                                print(key + "-->" + value)
                                para.runs[i].text = para.runs[i].text.replace(key, value)
    return doc


def clear_runs(runs):
    for i, run in enumerate(runs):
        if i > 0:
            run.clear()
    return runs


def debug_doc(doc):
    for i, sect in enumerate(doc.sections):
        for j, para in enumerate(sect.header.paragraphs):
            print(f'Sec.{i} Para.{j} : ', para.text, sep='')
            for k, run in enumerate(para.runs):
                print(f'Sec.{i} Para.{j} Run{k}: ', run.text, sep='')
    for i, para in enumerate(doc.paragraphs):
        print(f'Para.{i} : ', para.text, sep='')
        for j, run in enumerate(para.runs):
            print(f'Para.{i} Run{j}: ', run.text, sep='')
    for k, table in enumerate(doc.tables):
        for l, row in enumerate(table.rows):
            for m, cell in enumerate(row.cells):
                for i, para in enumerate(cell.paragraphs):
                    print(f'Table.{k} Row.{l} Cell{m} Para.{i} : ', para.text, sep='')
                    # for j, run in enumerate(para.runs):
                    #     print(f'Para.{i} Run{j}: ', run.text, sep='')


def replace_header(doc, prj):
    check_replace(doc.sections[0].header.paragraphs, '.*公司', prj.com_name)


def first_table(doc, prj):
    match = 0
    if doc.tables[0].rows[0].cells[0].paragraphs[0].text == '项目名称：':
        doc.tables[0].rows[0].cells[1].paragraphs[0].runs[0].text = prj.p_name
        clear_runs(doc.tables[0].rows[0].cells[1].paragraphs[0].runs)
        match = match + 1
    if doc.tables[0].rows[1].cells[0].paragraphs[0].text == '项目编号：':
        doc.tables[0].rows[1].cells[1].paragraphs[0].runs[0].text = prj.p_start[0:4] + 'RD' + prj.p_order
        clear_runs(doc.tables[0].rows[1].cells[1].paragraphs[0].runs)
        match = match + 1
    if doc.tables[0].rows[2].cells[0].paragraphs[0].text == '项目负责人：' and prj.p_owner != 'None':
        doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].text = prj.p_owner
        clear_runs(doc.tables[0].rows[2].cells[1].paragraphs[0].runs)
        match = match + 1
    if doc.tables[0].rows[3].cells[0].paragraphs[0].text == '项目周期：':
        doc.tables[0].rows[3].cells[1].paragraphs[0].runs[0].text = prj.p_start + '至' + prj.p_end
        clear_runs(doc.tables[0].rows[3].cells[1].paragraphs[0].runs)
        match = match + 1
    return match


def start_time(doc, prj):
    match = check_replace(doc.paragraphs, '申请立项时间：\d{4}[-/]\d{1,2}[-/]\d{1,2}', '申请立项时间：' + prj.p_start)
    return match


def second_table(doc, prj):
    match = 0
    if doc.tables[1].rows[0].cells[0].paragraphs[0].text == '项目立项名称':
        doc.tables[1].rows[0].cells[1].paragraphs[0].runs[0].text = prj.p_name
        clear_runs(doc.tables[1].rows[0].cells[1].paragraphs[0].runs)
        match = match + 1

    match = match + check_replace(doc.tables[1].rows[1].cells[1].paragraphs
                                  , '项目团队由(.*)人组成，项目实施周期为(.*)个月。'
                                  , '项目团队由' + prj.p_people + '人组成，项目实施周期为' + prj.p_cost + '个月。')
    match = match + check_replace(doc.tables[1].rows[6].cells[1].paragraphs
                                  , '\d{4}[-/]\d{1,2}[-/]\d{1,2}至\d{4}[-/]\d{1,2}[-/]\d{1,2}',
                                  prj.p_start + '至' + prj.p_end)
    match = match + check_replace(doc.tables[1].rows[7].cells[1].paragraphs
                                  , '项目总资金预算.*万元', '项目总资金预算' + prj.p_money + '万元')

    match = match + check_replace(doc.tables[1].rows[8].cells[1].paragraphs
                                  , '项目总人数：.*人', '项目总人数：' + prj.p_people + '人')
    if prj.p_owner != 'None':
        match = match + check_replace(doc.tables[1].rows[8].cells[1].paragraphs, '项目负责人：.*',
                                      '项目负责人：' + prj.p_owner)
    if prj.p_rnd != 'None':
        match = match + check_replace(doc.tables[1].rows[8].cells[1].paragraphs, '研发成员：.*', '研发成员：' + prj.p_rnd)
    match = match + check_replace(doc.tables[1].rows[9].cells[1].paragraphs, '\d{4}[-/]\d{1,2}[-/]\d{1,2}', prj.p_start)

    return match


def check_replace(paras, regex, dst):
    match = 0
    for i, para in enumerate(paras):
        result = re.search(regex, para.text)
        if result is not None:
            if result.group() != dst:
                print(result.group() + ' 被替换为 ' + dst)
                para.runs[0].text = re.sub(regex, dst,
                                           para.text)
                clear_runs(para.runs)
            match = match + 1
            break  # 只替换一次就够用
    return match


def third_table(doc, prj):
    match = 0
    if doc.tables[2].rows[0].cells[0].paragraphs[0].text == '项目名称':
        doc.tables[2].rows[0].cells[1].paragraphs[0].runs[0].text = prj.p_name
        clear_runs(doc.tables[1].rows[0].cells[1].paragraphs[0].runs)
        match = match + 1
    match = match + check_replace(doc.tables[2].rows[1].cells[1].paragraphs
                  , '\d{4}[-/]\d{1,2}[-/]\d{1,2}', prj.p_end)
    match = match + check_replace(doc.tables[2].rows[2].cells[1].paragraphs
                  , '\d{4}[-/]\d{1,2}[-/]\d{1,2}至\d{4}[-/]\d{1,2}[-/]\d{1,2}', prj.p_start + '至' + prj.p_end)

    if prj.p_owner != 'None':
        doc.tables[2].rows[3].cells[1].paragraphs[0].runs[0].text = prj.p_owner
        clear_runs(doc.tables[2].rows[3].cells[1].paragraphs[0].runs)
        match = match + 1
    return match


if __name__ == '__main__':
    init(autoreset=True)
    workdir = ''
    workdir_change = False
    config = ConfigParser()
    try:
        config.read('config.ini', encoding='UTF-8')
        workdir = config['config']['lasting']
    except:
        config.add_section('config')
        pass
    print('上次处理文件夹: ' + Fore.RED + workdir)
    yesno = input("直接回车继续处理。否则请输入新的路径：")
    if yesno != '':
        workdir_change = True
        workdir = yesno
        config['config']['lasting'] = yesno
    print('开始处理文件夹: ' + workdir)
    if not os.path.exists(workdir + '_bak'):
        print('''
        !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
        !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
        !!                                                           !!
        !! 该文件夹的文件会被自动修改，强烈建议备份文件，否则可能存在数据丢失风险  !!
        !!        备份文件夹以原名称加_bak扩展时不再提示该告警!              !!
        !!                                                           !!
        !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
        !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
        ''')
    input("输入回车开始执行修改文件,或者关闭该程序")

    for file_sum in os.listdir(workdir):
        if file_sum.endswith('立项报告汇总表.xlsx') and not file_sum.startswith('~$'):
            print(f"找到 {file_sum}, 开始处理.")
            break
    if not file_sum.endswith('立项报告汇总表.xlsx'):
        print("Error: 找不到 立项报告汇总表")
        sys.exit(-1)

    wb = load_workbook(workdir + '/' + file_sum, data_only=True)
    ws = wb.active
    if str(ws['A1'].value).find(u'公司') != -1:
        com_name = str(ws['A1'].value).split("公司")[0] + '公司'
    else:
        print("Error: 找不到 公司名")
        sys.exit(-2)

    max_row_num = ws.max_row
    rangeCell = ws[f'A3:P{max_row_num}']
    for r in rangeCell:
        if r[0].value is None:
            break
        project = Project()
        project.p_comname = com_name
        project.p_order = str(r[0].value).strip().zfill(2)
        project.p_name = str(r[1].value).strip()
        project.p_start = r[2].value.strftime('%Y-%m-%d')
        project.p_end = r[3].value.strftime('%Y-%m-%d')
        project.p_cost = str(r[5].value).strip()
        project.p_people = str(r[6].value).strip()  # 人数
        project.p_owner = str(r[7].value).strip()  # 项目负责人
        project.p_rnd = str(r[8].value).strip()  # 研发人员
        project.p_money = str(r[9].value).strip()  # 总预算

        try:
            doc_name = workdir + '/RD' + project.p_order + project.p_name + '.docx'
            document = Document(doc_name)
            # debug_doc(document)
            # TODO to be fixed
            # replace_header(document, project)
            first_table(document, project)
            start_time(document, project)
            second_table(document, project)
            third_table(document, project)
            document.save(doc_name)
        except PackageNotFoundError:
            print(Fore.RED + '打开文件错误：' + doc_name)

    #         document = replace_header(document, company)
    #         document = check_and_change(document, replace_dict)
    #         document = replace_tables(document, replace_dict)

    if workdir_change:
        with open('../../config.ini', 'w', encoding='utf-8') as file:
            config = ConfigParser()
            config.write(file)  # 数据写入配置文件
    input("处理完成.")
