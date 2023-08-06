# -*- coding: utf-8 -*-

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from datetime import datetime, date, time, timedelta

# 获取当前日期
nowTime = datetime.now()
now = nowTime.date()  # 2020-09-19
# 前一天
prevDayTime = nowTime + timedelta(days=-1)
prevDayDate = prevDayTime.date().strftime('%Y-%m-%d')

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
}

han_num = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
           "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
           "二十一", "二十二", "二十三", "二十四", "二十五", "二十六", "二十七", "二十八", "二十九", "三十"]

filename = "合规日报(" + prevDayTime.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日') + ").docx"

data_sources = [
    {"category": "监管动态", "sources": [
        {"title": "证监会要闻", "url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwdd/",
         "href_prefix_url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwdd/"},
        {"title": "新闻发布会", "url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwfbh/",
         "href_prefix_url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwfbh/"},
        {"title": "机关部门最新更新", "url": "http://www.csrc.gov.cn/pub/newsite/zxgx/jigbsdt/",
         "href_prefix_url": "http://www.csrc.gov.cn/pub/"},
        {"title": "派出机构最新更新", "url": "http://www.csrc.gov.cn/pub/newsite/zxgx/pcjgdt/",
         "href_prefix_url": "http://www.csrc.gov.cn/pub/"},
        {"title": "机构部-证券公司“白名单”", "url": "http://www.csrc.gov.cn/pub/newsite/zqjjjgjgb/zqgsbmd/",
         "href_prefix_url": "http://www.csrc.gov.cn/pub/newsite/zqjjjgjgb/zqgsbmd/"},
    ]}
]


# 解析网页
def parse_page(_url):
    response = requests.get(_url, headers=headers)
    text = response.content.decode("utf-8")
    return text


# 生产soup  pip install html5lib
def generate_soup(_text):
    _soup = BeautifulSoup(_text, "html5lib")
    return _soup


def get_soup(_url):
    return generate_soup(parse_page(_url))


def parse_url(_url):
    if _url.startswith("./"):
        return _url.replace("./", "")
    elif _url.startswith("../../../"):
        return _url.replace("../../../", "")
    else:
        return _url


def write_to_word(_category, titles, contents, document):
    for index, title in enumerate(titles):
        document.add_heading(han_num[index + 1] + "、" + title, level=2)
        # 添加段落
        p = document.add_paragraph(contents[index])


def parse_file(_category, _url, href_prefix_url, _document):
    soup = get_soup(_url)
    lis = []
    if len(soup.select("#myul li")) > 0:
        lis = lis + soup.select("#myul li")
    elif len(soup.select(".fl_list ul li")) > 0:
        lis = lis + soup.select(".fl_list ul li")
    titles = []
    hrefs = []
    contents = []
    for li in lis:
        time = li.find("span").get_text().strip()
        # # 如果发布时间不是昨天的，就放弃,继续下一次循环
        if prevDayDate != time:
            continue
        titles.append(li.find("a").get_text())
        hrefs.append(li.find("a")["href"])
    if len(hrefs) > 0:
        hrefs = list(map(lambda href: href_prefix_url + parse_url(href), hrefs))
    for index, title in enumerate(titles):
        print(index)
        print(title)
        print(hrefs[index])
        soup = get_soup(hrefs[index])
        content = []
        # if len(soup.select(".Custom_UnionStyle")) > 0:
        #     content = content + soup.select(".Custom_UnionStyle")
        if len(soup.select(".content p")) > 0:
            # nth-child(2+n) 第二个元素后面的元素
            content = soup.select(".content p")
        elif len(soup.select(".mainContainer .content")) > 0:
            # nth-child(2+n) 第二个元素后面的元素
            content = soup.select(".mainContainer .content")
        if len(content) > 0:
            # contents.append(content[0].get_text())
            str = ""
            for c in content:
                str = str + "\n".join(filter(None, c.get_text().split("\n"))) + "\n"
                # str = str + c.get_text() + "\n"
            contents.append(str)
        else:
            print(index, ":有页面没有获取到内容")
            contents.append("")
        print(contents[index])
    write_to_word(_category, titles, contents, _document)


if __name__ == "__main__":
    # 将数据保存到word中
    # 创建文档对象
    document = Document()
    # 全局设置下字体
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for index, data_source in enumerate(data_sources):
        category = data_source["category"]
        sources = data_source["sources"]
        # 添加标题
        document.add_heading("第" + han_num[index + 1] + "部分 " + category, level=1)
        for source in sources:
            # title = source["title"]
            url = source["url"]
            href_prefix_url = source["href_prefix_url"]
            parse_file(category, url, href_prefix_url, document)
    document.save("files/" + filename)
