# -*- coding: utf-8 -*-
from meutils.pipe import *
from meutils.request_utils.crawler import Crawler
from docx import Document
from docx.oxml.ns import qn
from datetime import datetime, date, time, timedelta

# 获取当前日期
nowTime = datetime.now()
now = nowTime.date()  # 2020-09-19
# 昨天
prevDayTime = nowTime + timedelta(days=-1)
prevDayDate = prevDayTime.date().strftime('%Y-%m-%d')

han_num = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
           "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
           "二十一", "二十二", "二十三", "二十四", "二十五", "二十六", "二十七", "二十八", "二十九", "三十"]
filePath = "files/"
filename = "合规日报(" + prevDayTime.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日') + ").docx"

data_sources = [
    {"category": "监管动态", "sources": [
        # {"name": "证监会要闻", "url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwdd/",
        #  "href_prefix_url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwdd/"},
        # {"name": "新闻发布会", "url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwfbh/",
        #  "href_prefix_url": "http://www.csrc.gov.cn/pub/newsite/zjhxwfb/xwfbh/"},
        # {"name": "机关部门最新更新", "url": "http://www.csrc.gov.cn/pub/newsite/zxgx/jigbsdt/",
        #  "href_prefix_url": "http://www.csrc.gov.cn/pub/"},
        {"name": "派出机构最新更新", "url": "http://www.csrc.gov.cn/pub/newsite/zxgx/pcjgdt/",
         "href_prefix_url": "http://www.csrc.gov.cn/pub/"},
        # {"name": "机构部-证券公司“白名单”", "url": "http://www.csrc.gov.cn/pub/newsite/zqjjjgjgb/zqgsbmd/",
        #  "href_prefix_url": "http://www.csrc.gov.cn/pub/newsite/zqjjjgjgb/zqgsbmd/"},
    ]}
]


def calc_href(_href):
    if _href.startswith("./"):
        return _href.replace("./", "")
    elif _href.startswith("../../../"):
        return _href.replace("../../../", "")
    else:
        return _href


# 获取网页的标题列表
def get_titles(_url):
    crawler = Crawler(_url)
    _titles = []
    _lis = crawler.xpath('//*[@id="myul"]//li') or crawler.xpath('//*[@class="fl_list"]//li')
    for _li in _lis:
        time = _li.xpath("span//text()")[0].strip()
        # 如果发布时间不是昨天的，就放弃,继续下一次循环
        if prevDayDate != time:
            continue
        title = _li.xpath("a//text()")[0]
        _titles.append(title)
    # 获取title对应的文本内容
    return _titles


def get_contents(_url, _prefix_url):
    crawler = Crawler(_url)
    # 获取title对应的a网址连接
    _hrefs = []
    _lis = crawler.xpath('//*[@id="myul"]//li') or crawler.xpath('//*[@class="fl_list"]//li')
    for _li in _lis:
        time = _li.xpath("span//text()")[0].strip()
        # 如果发布时间不是昨天的，就放弃,继续下一次循环
        if prevDayDate != time:
            continue
        href = _li.xpath("a//@href")[0]
        _hrefs.append(href)
    _href_urls = list(map(lambda _href: _prefix_url + calc_href(_href), _hrefs))
    _contents = []
    # 获取href链接的网址内容
    for _hrefs_url in _href_urls:
        crawler = Crawler(_hrefs_url)
        _content = crawler.xpath("""//*[@class="Custom_UnionStyle"]//text()""") or crawler.xpath(
            """//*[@class="content"]//p//text()""")
        _contents.append("".join(_content))  # 将内容列表合并成一个长字符串
    return _contents


if __name__ == "__main__":
    # 将数据保存到word中
    # 创建文档对象
    document = Document()
    # # 全局设置下字体
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # 写入内容
    for index, data_source in enumerate(data_sources):
        category = data_source["category"]
        sources = data_source["sources"]
        # 添加一级标题
        document.add_heading("第" + han_num[index + 1] + "部分 " + category, level=1)

        titles = []
        contents = []
        for source in sources:
            name = source["name"]
            url = source["url"]
            href_prefix_url = source["href_prefix_url"]
            # 获取这个网址下的标题列表titles
            titles = get_titles(url)
            contents = get_contents(url, href_prefix_url)

        # 写入内容
        for index, title in enumerate(titles):
            document.add_heading(han_num[index + 1] + "、" + title, level=2)
            # 添加段落
            p = document.add_paragraph(contents[index])

    document.save(filePath + filename)
